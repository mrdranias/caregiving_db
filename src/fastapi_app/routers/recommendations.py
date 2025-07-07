from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from sqlalchemy import text
from typing import List, Dict, Optional, Any
import logging
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown2
from database import get_db
from uuid import UUID as UUID_type
from models.adl_answers import ADLAnswers
from models.iadl_answers import IADLAnswers
from models.patient_history import PatientHistory
from models.risk import Risk
from models.hazards import Hazard

router = APIRouter(prefix="/recommendations", tags=["recommendations"])
logger = logging.getLogger(__name__)

@router.get("/by_patient/{patient_id}")
def get_recommendations(patient_id: str, db: Session = Depends(get_db)):
    """
    Get service recommendations for a patient based on their hazards and risk ratings.
    """
    try:
        try:
            uuid_obj = UUID_type(patient_id)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid patient_id format (must be UUID)")

        # First get hazards using same logic as hazards router
        hazards = []
        
        # --- ADL ---
        adl = db.query(ADLAnswers).filter(ADLAnswers.patient_id == uuid_obj).order_by(ADLAnswers.date_completed.desc()).first()
        if adl:
            adl_map = db.execute(text("SELECT adl_item, score_min, score_max, hazard_subclass_id, hazard_class_id FROM adl_item_hazard_map")).fetchall()
            adl_fields = [
                ("feeding", adl.feeding), ("bathing", adl.bathing), ("grooming", adl.grooming), ("dressing", adl.dressing),
                ("toilet_use", adl.toilet_use), ("transfers", adl.transfers)
            ]
            for item, score in adl_fields:
                if score is None:
                    continue
                for row in adl_map:
                    if row[0] == item and row[1] <= score <= row[2]:
                        hazard = {"type": "adl", "item": item, "score": score}
                        if row[3]:  # hazard_subclass_id
                            hazard["hazard_subclass_id"] = row[3]
                            hazard["hazard_code"] = row[3]
                        elif row[4]:  # hazard_class_id
                            hazard["hazard_class_id"] = row[4]
                            hazard["hazard_code"] = row[4]
                        hazards.append(hazard)

        # --- IADL ---
        iadl = db.query(IADLAnswers).filter(IADLAnswers.patient_id == uuid_obj).order_by(IADLAnswers.date_completed.desc()).first()
        if iadl:
            iadl_map = db.execute(text("SELECT iadl_item, score_min, score_max, hazard_subclass_id, hazard_class_id FROM iadl_item_hazard_map")).fetchall()
            iadl_fields = [
                ("shopping", iadl.shopping), ("food_preparation", iadl.food_preparation), ("housekeeping", iadl.housekeeping),
                ("transportation", iadl.transportation), ("medication", iadl.medication), ("finances", iadl.finances)
            ]
            for item, score in iadl_fields:
                if score is None:
                    continue
                for row in iadl_map:
                    if row[0] == item and row[1] <= score <= row[2]:
                        hazard = {"type": "iadl", "item": item, "score": score}
                        if row[3]:  # hazard_subclass_id
                            hazard["hazard_subclass_id"] = row[3]
                            hazard["hazard_code"] = row[3]
                        elif row[4]:  # hazard_class_id
                            hazard["hazard_class_id"] = row[4]
                            hazard["hazard_code"] = row[4]
                        hazards.append(hazard)

        # --- Patient History (Dx, Sx, Rx) ---
        history = db.query(PatientHistory).filter(PatientHistory.patient_id == uuid_obj).order_by(PatientHistory.created_at.desc()).first()
        if history:
            # Sx
            if history.sx_codes:
                sx_map = db.execute(text("SELECT sx_code, hazard_subclass_id, hazard_class_id FROM sx_code_hazard_map")).fetchall()
                for code in history.sx_codes:
                    for row in sx_map:
                        if row[0] == code:
                            hazard = {"type": "sx", "code": code}
                            if row[1]:  # hazard_subclass_id
                                hazard["hazard_subclass_id"] = row[1]
                                hazard["hazard_code"] = row[1]
                            elif row[2]:  # hazard_class_id
                                hazard["hazard_class_id"] = row[2]
                                hazard["hazard_code"] = row[2]
                            hazards.append(hazard)
            
            # Dx
            if history.dx_codes:
                dx_map = db.execute(text("SELECT dx_code, hazard_subclass_id, hazard_class_id FROM dx_code_hazard_map")).fetchall()
                for code in history.dx_codes:
                    for row in dx_map:
                        if row[0] == code:
                            hazard = {"type": "dx", "code": code}
                            if row[1]:  # hazard_subclass_id
                                hazard["hazard_subclass_id"] = row[1]
                                hazard["hazard_code"] = row[1]
                            elif row[2]:  # hazard_class_id
                                hazard["hazard_class_id"] = row[2]
                                hazard["hazard_code"] = row[2]
                            hazards.append(hazard)

        # --- Social Hazards (from existing social_risks table) ---
        # Use existing social risks that are already processed and stored
        from models.social_risk import SocialRisk
        
        social_risks = db.query(SocialRisk).filter(SocialRisk.patient_id == uuid_obj).all()
        for social_risk in social_risks:
            if social_risk.risk_score and social_risk.risk_score > 0:  # Only include active risks
                hazard = {
                    "type": "social",
                    "hazard_code": social_risk.social_hazard_code,
                    "hazard_subclass_id": social_risk.social_hazard_code,
                    "social_hazard_type": social_risk.social_hazard_type,
                    "risk_score": social_risk.risk_score
                }
                hazards.append(hazard)

        if not hazards:
            return {"service_categories": [], "total_services": 0}

        # Now get service mappings and risks for each hazard
        service_mappings = {}
        risks_by_hazard = {}
        
        # Get all clinical risk data for this patient (from risks table)
        risks_query = db.query(Risk, Hazard).join(Hazard, Risk.hazard_id == Hazard.hazard_id).filter(Risk.patient_id == uuid_obj).all()
        for risk, hazard in risks_query:
            # Index by the hazard_type (which matches hazard_code)
            risks_by_hazard[hazard.hazard_type] = {
                "risk_id": str(risk.risk_id),
                "severity": risk.severity,
                "likelihood": risk.likelihood,
                "risk_score": risk.severity * risk.likelihood if risk.severity and risk.likelihood else 0,
                "notes": risk.notes,
                "hazard_type": "clinical"  # Explicitly mark as clinical
            }
        
        # Get all social risk data for this patient (from social_risks table)
        social_risks_query = db.query(SocialRisk).filter(SocialRisk.patient_id == uuid_obj).all()
        for social_risk in social_risks_query:
            if social_risk.risk_score and social_risk.risk_score > 0:  # Only include active risks
                # Index by social_hazard_code to match hazard lookup
                risks_by_hazard[social_risk.social_hazard_code] = {
                    "risk_id": str(social_risk.social_risk_id),
                    "severity": None,  # Social risks use direct risk_score
                    "likelihood": None,
                    "risk_score": social_risk.risk_score,
                    "notes": social_risk.notes or "",
                    "hazard_type": "social"  # Explicitly mark as social
                }

        # Get service mappings for each hazard
        for hazard in hazards:
            hazard_code = hazard.get("hazard_code")
            hazard_type = hazard.get("type", "clinical")
            if not hazard_code:
                continue
                
            service_rows = []
            
            # Handle social hazards (use SDOH mitigation tables)
            if hazard_type == "social":
                # Try SDOH mitigation subclass mappings first
                sdoh_subclass_query = text("""
                    SELECT DISTINCT
                        sdm.mitigation_subclass_id as service_subclass_id,
                        ms.label as service_subclass_label,
                        ms.description as service_subclass_description,
                        ms.parent_class_id as service_class_id,
                        mc.label as service_class_label,
                        mc.description as service_class_description
                    FROM sdoh_mitigation_map sdm
                    LEFT JOIN sdoh_mitigation_subclasses ms ON sdm.mitigation_subclass_id = ms.subclass_id
                    LEFT JOIN sdoh_mitigations mc ON ms.parent_class_id = mc.class_id
                    WHERE sdm.social_hazard_subclass_id = :hazard_code
                    AND sdm.mitigation_subclass_id IS NOT NULL
                """)
                
                service_rows = db.execute(sdoh_subclass_query, {"hazard_code": hazard_code}).fetchall()
                
                # If no subclass mappings found, try parent SDOH mappings
                if not service_rows:
                    # Get parent social hazard class
                    social_parent_query = text("""
                        SELECT shs.parent_class_id 
                        FROM social_hazards_subclasses shs 
                        WHERE shs.subclass_id = :hazard_code
                    """)
                    social_parent_result = db.execute(social_parent_query, {"hazard_code": hazard_code}).fetchone()
                    
                    if social_parent_result and social_parent_result[0]:
                        parent_social_hazard_class = social_parent_result[0]
                        
                        # Get mitigation classes mapped to this parent social hazard class
                        parent_sdoh_query = text("""
                            SELECT DISTINCT
                                NULL as service_subclass_id,
                                'General SDOH Services' as service_subclass_label,
                                'General services for this social hazard category' as service_subclass_description,
                                psm.mitigation_class_id as service_class_id,
                                mc.label as service_class_label,
                                mc.description as service_class_description
                            FROM parent_sdoh_mitigation_map psm
                            LEFT JOIN sdoh_mitigations mc ON psm.mitigation_class_id = mc.class_id
                            WHERE psm.social_hazard_class_id = :parent_hazard_class
                        """)
                        
                        service_rows = db.execute(parent_sdoh_query, {"parent_hazard_class": parent_social_hazard_class}).fetchall()
            else:
                # Handle clinical hazards (use existing service tables)
                # First try to get service mappings from subclass level (hazard_service_map)
                subclass_query = text("""
                    SELECT DISTINCT
                        hsm.service_subclass_id,
                        ss.label as service_subclass_label,
                        ss.description as service_subclass_description,
                        ss.parent_class_id as service_class_id,
                        sc.label as service_class_label,
                        sc.description as service_class_description
                    FROM hazard_service_map hsm
                    LEFT JOIN service_subclasses ss ON hsm.service_subclass_id = ss.subclass_id
                    LEFT JOIN service_classes sc ON ss.parent_class_id = sc.class_id
                    WHERE hsm.hazard_subclass_id = :hazard_code
                    AND hsm.service_subclass_id IS NOT NULL
                """)
                
                service_rows = db.execute(subclass_query, {"hazard_code": hazard_code}).fetchall()
            
            # If no subclass mappings found, fallback to parent class mappings
            if not service_rows:
                # Get the parent hazard class for this hazard
                hazard_parent_query = text("""
                    SELECT hs.parent_class_id 
                    FROM hazard_subclasses hs 
                    WHERE hs.subclass_id = :hazard_code
                """)
                hazard_parent_result = db.execute(hazard_parent_query, {"hazard_code": hazard_code}).fetchone()
                
                if hazard_parent_result and hazard_parent_result[0]:
                    parent_hazard_class = hazard_parent_result[0]
                    
                    # Get service classes mapped to this parent hazard class
                    parent_query = text("""
                        SELECT DISTINCT
                            NULL as service_subclass_id,
                            'General Services' as service_subclass_label,
                            'General services for this hazard category' as service_subclass_description,
                            psm.service_class_id,
                            sc.label as service_class_label,
                            sc.description as service_class_description
                        FROM parent_hazard_service_map psm
                        LEFT JOIN service_classes sc ON psm.service_class_id = sc.class_id
                        WHERE psm.hazard_class_id = :hazard_class_id
                    """)
                    
                    service_rows = db.execute(parent_query, {"hazard_class_id": parent_hazard_class}).fetchall()
            
            for row in service_rows:
                service_class_id = row[3] or "Uncategorized"
                service_class_label = row[4] or "Uncategorized Services"
                
                if service_class_id not in service_mappings:
                    service_mappings[service_class_id] = {
                        "service_class_id": service_class_id,
                        "service_class_label": service_class_label,
                        "service_class_description": row[5] or "",
                        "services": []
                    }
                
                # Get risk data for this hazard using hazard_code
                risk_data = risks_by_hazard.get(hazard_code, {})
                risk_score = risk_data.get("risk_score", 0)
                
                # Determine priority based on risk score
                if risk_score >= 20:
                    priority = "High"
                elif risk_score >= 10:
                    priority = "Medium"
                elif risk_score > 0:
                    priority = "Low"
                else:
                    priority = "Info"
                
                service_info = {
                    "service_subclass_id": row[0],
                    "service_subclass_label": row[1] or "Unknown Service",
                    "service_subclass_description": row[2] or "",
                    "linked_hazards": [{
                        "hazard_code": hazard_code,
                        "hazard_type": hazard.get("type", "") or risk_data.get("hazard_type", ""),
                        "hazard_item": hazard.get("item", ""),
                        "hazard_diagnosis_code": hazard.get("code", ""),
                        "risk_score": risk_score,
                        "severity": risk_data.get("severity"),
                        "likelihood": risk_data.get("likelihood"),
                        "notes": risk_data.get("notes", "")
                    }],
                    "priority": priority,
                    "max_risk_score": risk_score
                }
                
                # Check if service already exists and merge hazards
                existing_service = None
                for service in service_mappings[service_class_id]["services"]:
                    if service["service_subclass_id"] == row[0]:
                        existing_service = service
                        break
                
                if existing_service:
                    existing_service["linked_hazards"].extend(service_info["linked_hazards"])
                    existing_service["max_risk_score"] = max(existing_service["max_risk_score"], risk_score)
                    # Update priority if this risk is higher
                    if risk_score > existing_service.get("max_risk_score", 0):
                        existing_service["priority"] = priority
                else:
                    service_mappings[service_class_id]["services"].append(service_info)

        # Convert to list and sort by priority
        service_categories = list(service_mappings.values())
        for category in service_categories:
            category["services"].sort(key=lambda x: (-x["max_risk_score"], x["service_subclass_label"]))
            
            # Calculate total_risk_score and hazard_count for this category
            total_risk_score = sum(service["max_risk_score"] for service in category["services"])
            hazard_count = sum(len(service["linked_hazards"]) for service in category["services"])
            
            category["total_risk_score"] = total_risk_score
            category["hazard_count"] = hazard_count
        
        service_categories.sort(key=lambda x: x["service_class_label"])
        
        total_services = sum(len(cat["services"]) for cat in service_categories)
        
        return {
            "patient_id": patient_id,
            "service_categories": service_categories,
            "total_services": total_services,
            "total_service_categories": len(service_categories)
        }
        
    except Exception as e:
        logging.error(f"Error fetching recommendations for patient {patient_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching recommendations: {str(e)}")


def filter_selected_recommendations(patient_id: str, recommendations: dict, db: Session) -> dict:
    """
    Filter recommendations to only include services that are marked as selected in recommendation_settings.
    """
    try:
        uuid_obj = UUID_type(patient_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid patient_id format (must be UUID)")
    
    # Get selected services from recommendation_settings
    selected_services_query = text("""
        SELECT hazard_code, service_description, service_category
        FROM recommendation_settings 
        WHERE patient_id = :patient_id AND selected = TRUE
    """)
    
    selected_services = db.execute(selected_services_query, {"patient_id": str(uuid_obj)}).fetchall()
    
    # Create a set of selected service identifiers for fast lookup
    selected_service_keys = set()
    for service in selected_services:
        # Create unique key combining hazard_code and service_description
        key = f"{service.hazard_code}|{service.service_description}"
        selected_service_keys.add(key)
    
    # Filter the recommendations
    filtered_recommendations = {
        "patient_id": patient_id,
        "service_categories": [],
        "total_services": 0,
        "total_service_categories": 0,
        "total_contractors": recommendations.get("total_contractors", 0)
    }
    
    # Go through each service category and filter services
    for category in recommendations.get("service_categories", []):
        filtered_services = []
        
        for service in category.get("services", []):
            # Check if any of the linked hazards for this service are selected
            service_selected = False
            for hazard in service.get("linked_hazards", []):
                hazard_code = hazard.get("hazard_code", "")
                service_description = service.get("service_subclass_label", "")
                key = f"{hazard_code}|{service_description}"
                
                if key in selected_service_keys:
                    service_selected = True
                    break
            
            if service_selected:
                filtered_services.append(service)
        
        # Only include categories that have selected services
        if filtered_services:
            filtered_category = category.copy()
            filtered_category["services"] = filtered_services
            filtered_recommendations["service_categories"].append(filtered_category)
    
    filtered_recommendations["total_services"] = sum(
        len(cat["services"]) for cat in filtered_recommendations["service_categories"]
    )
    filtered_recommendations["total_service_categories"] = len(filtered_recommendations["service_categories"])
    
    return filtered_recommendations

@router.post("/generate_report/{patient_id}")
async def generate_recommendation_report(patient_id: str, db: Session = Depends(get_db)):
    """
    Generate and save a comprehensive recommendation report for a patient.
    Stores the report in the recommendation_report table.
    """
    try:
        logger.info(f"Starting generate_recommendation_report for patient {patient_id}")
        
        # Get recommendations data
        recommendations = get_recommendations(patient_id, db)
        logger.info(f"Retrieved recommendations data with {len(recommendations.get('service_categories', []))} categories")
        
        # Filter recommendations to only include selected services
        selected_recommendations = filter_selected_recommendations(patient_id, recommendations, db)
        logger.info(f"Filtered to {len(selected_recommendations.get('service_categories', []))} categories with selected services")
        
        if not selected_recommendations["service_categories"]:
            raise HTTPException(status_code=404, detail="No selected service recommendations available for report generation")
        
        # Generate report content
        logger.info("Generating report content...")
        try:
            report_content = generate_report_content(selected_recommendations)
            logger.info(f"Generated report content of length {len(report_content)}")
        except Exception as e:
            logger.error(f"Error in generate_report_content: {e}")
            raise HTTPException(status_code=500, detail=f"Error generating report content: {str(e)}")
        
        # Save to database
        logger.info(f"Starting report generation for patient {patient_id}")
        insert_query = text("""
        INSERT INTO recommendation_report (patient_id, risks, services, contractors, costs, content)
        VALUES (:patient_id, :risks, :services, :contractors, :costs, :content)
        RETURNING report_id, generated_on
        """)
        
        report_data = {
            "patient_id": patient_id,
            "risks": json.dumps(selected_recommendations),  # Store filtered recommendations as proper JSON
            "services": json.dumps([cat["services"] for cat in selected_recommendations["service_categories"]]),
            "contractors": json.dumps({}),  # Placeholder for contractors (future feature)
            "costs": json.dumps({}),  # Placeholder for costs (future feature)
            "content": report_content
        }
        
        logger.info(f"Executing insert query with data keys: {list(report_data.keys())}")
        result = db.execute(insert_query, report_data)
        
        report_info = result.fetchone()
        logger.info(f"Insert result: {report_info}")
        
        db.commit()
        logger.info(f"Transaction committed successfully")
        
        # Generate docx file
        logger.info("Creating docx file...")
        try:
            docx_filepath = create_docx_from_markdown(report_content, patient_id)
            logger.info(f"Created docx file at: {docx_filepath}")
            
            # File is now saved locally, so just use the path as-is
            relative_path = docx_filepath
            
        except Exception as e:
            logger.error(f"Error creating docx file: {e}")
            docx_filepath = None
            relative_path = None
        
        logger.info(f"Generated recommendation report {report_info.report_id} for patient {patient_id}")
        
        return {
            "report_id": str(report_info.report_id),
            "patient_id": patient_id,
            "generated_on": str(report_info.generated_on),
            "content": report_content,
            "recommendations": recommendations,
            "docx_file_path": docx_filepath,
            "docx_download_path": relative_path
        }
        
    except Exception as e:
        logger.error(f"Error generating report for patient {patient_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")


def generate_report_content(recommendations: Dict[str, Any]) -> str:
    """Generate human-readable report content from recommendations data."""
    content = f"# Service Recommendation Report\n\n"
    content += f"**Patient ID:** {recommendations['patient_id']}\n"
    content += f"**Total Service Categories:** {recommendations['total_service_categories']}\n"
    content += f"**Total Recommended Services:** {recommendations['total_services']}\n\n"
    
    for i, category in enumerate(recommendations["service_categories"], 1):
        content += f"## {i}. {category['service_class_label']}\n"
        content += f"*{category['service_class_description']}*\n"
        content += f"**Priority Score:** {category['total_risk_score']:.1f} (based on {category['hazard_count']} hazards)\n\n"
        
        for j, service in enumerate(category["services"], 1):
            content += f"### {i}.{j} {service['service_subclass_label']}\n"
            content += f"{service['service_subclass_description']}\n"
            content += f"**Priority Score:** {service['max_risk_score']:.1f}\n\n"
            
            content += "**Linked Hazards:**\n"
            for hazard in service["linked_hazards"]:
                hazard_desc = hazard["hazard_item"] or hazard["hazard_diagnosis_code"] or hazard["hazard_type"]
                risk_info = ""
                if hazard["risk_score"]:
                    risk_info = f" (Risk Score: {hazard['risk_score']:.1f})"
                hazard_id = hazard.get('hazard_subclass_id') or hazard.get('hazard_class_id', 'Unknown')
                content += f"- {hazard_id}: {hazard_desc}{risk_info}\n"
            content += "\n"
    
    return content


def create_docx_from_markdown(markdown_content: str, patient_id: str) -> str:
    """Convert markdown content to a formatted docx file and return the file path."""
    # Create document
    doc = Document()
    
    # Add header
    header_para = doc.sections[0].header.paragraphs[0]
    header_para.text = f"Care Management Report - Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    
    # Parse markdown content by lines
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Handle bold text and regular paragraphs
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif '**' in line:
            # Handle inline bold formatting
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    p.add_run(part)
                else:
                    run = p.add_run(part)
                    run.bold = True
        
        # Handle bullet points
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # Handle italic text
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True
        
        # Regular paragraphs
        else:
            doc.add_paragraph(line)
    
    # Create filename and save
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    # Use only first 6 characters of patient_id for shorter filename
    short_patient_id = patient_id[:6] if len(patient_id) >= 6 else patient_id
    filename = f"care_plan_{short_patient_id}_{timestamp}.docx"
    
    # Save to local reports directory (accessible outside container)
    reports_dir = "reports"  # Local reports directory
    os.makedirs(reports_dir, exist_ok=True)
    
    filepath = os.path.join(reports_dir, filename)
    doc.save(filepath)
    
    return filepath


# Pydantic models for save endpoint
from pydantic import BaseModel
from typing import List

class RecommendationSaveRequest(BaseModel):
    patient_id: str
    hazard_code: str
    service_description: str
    service_category: str
    frequency: str
    estimated_cost: float
    provider: str
    priority: str
    notes: str
    selected: bool = True  # Default to True for backwards compatibility

class RecommendationsBatchSaveRequest(BaseModel):
    patient_id: str
    recommendations: List[RecommendationSaveRequest]

@router.post("/save")
def save_recommendations(request: RecommendationsBatchSaveRequest, db: Session = Depends(get_db)):
    """
    Save user-entered recommendation settings (frequency, cost, provider) for a patient.
    """
    try:
        # Validate patient_id format
        try:
            uuid_obj = UUID_type(request.patient_id)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid patient_id format (must be UUID)")
        
        saved_count = 0
        
        # Create or update recommendation records
        for rec in request.recommendations:
            # Check if a recommendation for this patient/hazard/service already exists
            existing_query = text("""
                SELECT rec_id FROM recommendation_settings 
                WHERE patient_id = :patient_id 
                AND hazard_code = :hazard_code 
                AND service_description = :service_description
            """)
            
            existing = db.execute(existing_query, {
                "patient_id": str(uuid_obj),
                "hazard_code": rec.hazard_code,
                "service_description": rec.service_description
            }).fetchone()
            
            if existing:
                # Update existing record
                update_query = text("""
                    UPDATE recommendation_settings 
                    SET service_category = :service_category,
                        frequency = :frequency,
                        estimated_cost = :estimated_cost,
                        provider = :provider,
                        priority = :priority,
                        notes = :notes,
                        selected = :selected,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE rec_id = :rec_id
                """)
                
                db.execute(update_query, {
                    "rec_id": existing.rec_id,
                    "service_category": rec.service_category,
                    "frequency": rec.frequency,
                    "estimated_cost": rec.estimated_cost,
                    "provider": rec.provider,
                    "priority": rec.priority,
                    "notes": rec.notes,
                    "selected": rec.selected
                })
            else:
                # Insert new record
                insert_query = text("""
                    INSERT INTO recommendation_settings 
                    (patient_id, hazard_code, service_description, service_category, 
                     frequency, estimated_cost, provider, priority, notes, selected)
                    VALUES (:patient_id, :hazard_code, :service_description, :service_category,
                            :frequency, :estimated_cost, :provider, :priority, :notes, :selected)
                """)
                
                db.execute(insert_query, {
                    "patient_id": str(uuid_obj),
                    "hazard_code": rec.hazard_code,
                    "service_description": rec.service_description,
                    "service_category": rec.service_category,
                    "frequency": rec.frequency,
                    "estimated_cost": rec.estimated_cost,
                    "provider": rec.provider,
                    "priority": rec.priority,
                    "notes": rec.notes,
                    "selected": rec.selected
                })
            
            saved_count += 1
        
        db.commit()
        
        logger.info(f"Saved {saved_count} recommendation settings for patient {request.patient_id}")
        
        return {
            "message": f"Successfully saved {saved_count} recommendation settings",
            "patient_id": request.patient_id,
            "saved_count": saved_count
        }
        
    except Exception as e:
        db.rollback()
        logger.error(f"Error saving recommendations for patient {request.patient_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error saving recommendations: {str(e)}")
